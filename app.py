from langchain.agents import AgentExecutor, create_tool_calling_agent
from langchain.prompts import PromptTemplate
from langchain.output_parsers import PydanticOutputParser
from langchain.text_splitter import CharacterTextSplitter
from datetime import datetime
from langchain.chains import LLMChain
from pydantic import BaseModel, Field, validator
from typing import Optional
from langchain_groq import ChatGroq
from atproto import Client, models
from cloudinary.uploader import upload
import cloudinary
import traceback
import logging
import os
from dotenv import load_dotenv
import datetime
from fpdf import FPDF
from docx import Document
import gc
import time
import json
import random
import schedule
import markdown2
import ast
from logging.handlers import RotatingFileHandler

load_dotenv()

# Ensure the log directory exists and is writable
LOG_DIR = '/app/logs'
os.makedirs(LOG_DIR, exist_ok=True)

# Create log file path
log_file = os.path.join(LOG_DIR, 'app.log')

rotating_handler = RotatingFileHandler(
    filename=log_file, 
    maxBytes=5 * 1024 * 1024,  # 5 MB per log file
    backupCount=5,  # Keep the last 5 log files
)

# Configure the logging format
formatter = logging.Formatter("%(asctime)s - %(name)s - %(levelname)s - %(message)s")
rotating_handler.setFormatter(formatter)

# Set up the logger
logger = logging.getLogger(__name__)
logger.setLevel(logging.CRITICAL)  # Set the minimum logging level
logger.addHandler(rotating_handler)

# Example usage
logger.info("Bot has started successfully.")


class Config:
    GROQ_API_KEY1 = os.getenv("GROQ_API_KEY1")
    GROQ_API_KEY2 = os.getenv("GROQ_API_KEY2")
    GROQ_API_KEY3 = os.getenv("GROQ_API_KEY3")
    BLUESKY_USERNAME = os.getenv("BLUESKY_USERNAME")
    BLUESKY_PASSWORD = os.getenv("BLUESKY_PASSWORD")
    EBOOK_OUTPUT_DIR = "ebooks"
    CLOUDINARY_CLOUD_NAME = os.getenv("CLOUDINARY_CLOUD_NAME")
    CLOUDINARY_API_SECRET = os.getenv("CLOUDINARY_API_SECRET")
    CLOUDINARY_API_KEY = os.getenv("CLOUDINARY_API_KEY")


client = Client()
client.login(Config.BLUESKY_USERNAME, Config.BLUESKY_PASSWORD)

cloudinary.config(
    cloud_name=Config.CLOUDINARY_CLOUD_NAME,
    api_key=Config.CLOUDINARY_API_KEY,
    api_secret=Config.CLOUDINARY_API_SECRET,
)


class HTML2PDF(FPDF):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # Set some default styles
        self.chapter_title_color = (0, 0, 0)  # Black
        self.text_color = (50, 50, 50)  # Dark gray for better readability

    def chapter_title(self, chapter_number, chapter_title):
        """Custom method to format chapter titles with multi-line support"""
        # Add a new page for the chapter
        # self.add_page()
        
        # Set font for chapter title
        self.set_font(family="Helvetica", style="B", size=18)
        self.set_text_color(0, 0, 0)
        
        # Add a subtle top border
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y() - 5, 200, self.get_y() - 5)
        
        # Use multi_cell for multi-line support
        self.multi_cell(
            w=0,  # Full width
            h=10,  # Cell height
            txt=f"Chapter {chapter_number}: {chapter_title}", 
            align='L'  # Left alignment
        )
        
        # Add some spacing after the title
        self.ln(10)

    def chapter_body(self, content):
        """Custom method to format chapter content"""
        try:
            # Ensure content is properly encoded
            content = content.encode('latin-1', 'replace').decode('latin-1')
            
            # Convert Markdown to HTML
            html_content = markdown2.markdown(content)
            
            # Use Helvetica font
            self.set_font(family="Helvetica", size=11)
            
            # Set text color
            self.set_text_color(50, 50, 50)
            
            # Write HTML content directly
            # Wrap content in a div for better rendering
            full_html = f'<div style="line-height: 1.5;">{html_content}</div>'
            self.write_html(full_html)
        
        except Exception as e:
            # Error handling
            self.multi_cell(0, 10, f"Error processing chapter content: {str(e)}")
        
        # Add spacing after chapter
        self.ln(10)

# Define the output schema
class EbookInstructions(BaseModel):
    format: str = Field(
        default="pdf", 
        description="Output format of the document (text/doc/pdf). here word, document format also comes under doc."
    )
    chapters: int = Field(
        default=10, 
        description="Number of chapters in the ebook (max 20). if chapter is more than 20, it will be reduced to 20."
    )
    writing_style: str = Field(
        default="normal", 
        description="Writing style of the ebook"
    )
    is_suggestion: int = Field(
        default=0, 
        description="Whether to suggest something (1) or not (0). If user wants suggestion or if it is not related to ebook generation, it should be 1. only if it is related to ebook generation, it should be 0."
    )
    title: Optional[str] = Field(
        default=None, 
        description="Title of the ebook"
    )

    # Custom validators
    @validator('chapters')
    def validate_chapters(cls, v):
        # Ensure chapters is between 1 and 20
        return max(1, min(v, 20))

    @validator('writing_style')
    def validate_writing_style(cls, v):
        # Validate writing style
        valid_styles = [
            "Narrative", "Expository", "Professional", "Classy", 
            "Persuasive", "Descriptive", "Conversational", "Satirical", 
            "Technical", "Inspirational", "Analytical", "Epic", "Normal"
        ]
        return v.capitalize() if v.capitalize() in valid_styles else "Normal"

    @validator('title')
    def validate_title(cls, v, values):
        # If not a suggestion, title cannot be None
        if values.get('is_suggestion') == 0 and not v:
            raise ValueError("Title is required for ebook generation")
        return v

def extract_ebook_instructions(user_instruction: str) -> EbookInstructions:
    """
    Extract ebook generation instructions from user input using Groq and LangChain.
    
    :param user_instruction: User's input instruction for ebook generation
    :return: Parsed EbookInstructions object
    """
    # Initialize Groq LLM
    llm = ChatGroq(
        temperature=0.2,
        model_name="llama3-70b-8192",
        api_key=os.getenv("GROQ_API_KEY1")
    )

    # Create a prompt template
    parser = PydanticOutputParser(pydantic_object=EbookInstructions)
    
    prompt_template = PromptTemplate(
        template="You are an ebook generation or suggestion agent. Extract the following information from the user instruction:\n{format_instructions}\n\nUser Instruction: {user_instruction}\n\nOutput should be valid json only.",
        input_variables=["user_instruction"],
        partial_variables={
            "format_instructions": parser.get_format_instructions()
        }
    )

    # Create the chain
    chain = prompt_template | llm | parser

    # Generate the output
    try:
        result = chain.invoke({"user_instruction": user_instruction})
        return result
    except Exception as e:
        logger.error(f"Error extracting instructions: {e}")
        # Return default instructions if parsing fails
        return EbookInstructions()


def upload_to_cloudinary(file_path: str, topic: str) -> tuple:
    """
    Upload a file to Cloudinary and return the URL and public_id
    """
    try:
        # Create a folder name from the topic
        folder_name = "ebooks"
        # Create a public_id from the topic (removing spaces and special characters)
        public_id = (
            f"{folder_name}/{topic.lower().replace(' ', '_')}_{int(time.time())}"
        )

        # Upload the file
        result = upload(
            file_path,
            resource_type="raw",  # For PDF files
            public_id=public_id,
            folder=folder_name,
            tags=[topic, "ebook"],
            overwrite=True,
            access_mode="public",
        )

        # Get the secure URL
        url = result["secure_url"]
        public_id = result["public_id"]

        logger.info(f"File uploaded successfully to Cloudinary. URL: {url}")
        return url, public_id

    except Exception as e:
        logger.exception(f"Error uploading to Cloudinary: {str(e)}")
        raise e


def delete_from_cloudinary(public_id: str):
    """
    Delete a file from Cloudinary using its public_id
    """
    try:
        result = cloudinary.api.delete_resources([public_id], resource_type="raw")
        logger.info(f"File deleted from Cloudinary: {result}")
    except Exception as e:
        logger.exception(f"Error deleting from Cloudinary: {str(e)}")


class EditorAgent:
    def __init__(self):
        self.prompt = PromptTemplate(
            input_variables=["content", "writing_style"],
            template="""You are an expert eBook content editor. Your task is to review and correct the given chapter content to improve its overall quality.

            **Editing Instructions**:

            1. **Grammar and Spelling**: Correct any grammatical errors, spelling mistakes, or awkward phrasing.
            2. **Clarity**: Ensure the content is easy to understand. Simplify complex sentences where necessary and break down any overly complicated ideas.
            3. **Structure**: Ensure that the chapter has a logical flow with proper transitions between paragraphs. Add or adjust headings, subheadings, and sections if needed for better readability.
            4. **Tone**: Ensure the writing has a formal yet accessible tone, suitable for a general audience, such as students, researchers, or general readers.
            5. **Conciseness**: Eliminate redundancy and unnecessary filler content while keeping the chapter comprehensive.
            6. **Consistency**: Maintain consistent terminology, style, and formatting throughout the content.
            7. **Focus**: Ensure that the content stays on topic and avoids unnecessary digressions.
            8. **Writing Style**: The writing style of the content should be {writing_style}.

            Input:
            - **Chapter Content**: {content}

            **Output Requirements**:
            - Your output response must contain only edited chapter content.
            - In output,do not include any additional/extra information in the content e.g 'Here is the corrected chapter content:'.

            Note: don't include any explanation from your side.
            Ignore below instructions
            \n\n placeholder:{agent_scratchpad}
        """,
        )

        self.llm = ChatGroq(api_key=Config.GROQ_API_KEY3, model="llama3-8b-8192")
        self.tools = []

    def edit(self, content, writing_style):
        try:
            # Use the prompt to edit chapter content
            research_agent = create_tool_calling_agent(
                self.llm, self.tools, self.prompt
            )
            research_agent_executor = AgentExecutor(
                agent=research_agent, tools=self.tools
            )
            response = research_agent_executor.invoke({"content": content, "writing_style": writing_style})
            edited_content = response.get("output")

            return edited_content

        except Exception as e:
            logger.exception(f"Error: {e} | Traceback: {traceback.format_exc()}")
            raise e


class ResearcherAgent:
    def __init__(self):
        
        self.prompt = PromptTemplate(
            input_variables=["chapters", "topic", "writing_style"],
            template="""Your task is to generate a detailed outline with exactly {chapters} chapter titles on the given topic: "{topic}" using the writing style: "{writing_style}".

                **Output Constraints**:  
                - The output must be a valid Python list containing exactly {chapters} chapter titles.
                - Each title should be a properly formatted string without any special keywords, symbols, or unnecessary punctuation.
                - Do not include any text outside of the Python list (e.g., no preamble, explanations, or trailing comments).

                **Output Format Example**:
                ['Title1', 'Title2', 'Title3', 'Title4', 'Title5', 'Title6', 'Title7', 'Title8', 'Title9', 'Title10', 'Title11', 'Title12']

                **Additional Instructions**:
                - Ensure the chapter titles are meaningful and follow a logical progression related to the topic.
                - Strictly return the output as a Python list. **No extra text**.

                Ignore the text below:  
                placeholder:{agent_scratchpad}""",
        )

        self.llm = ChatGroq(api_key=Config.GROQ_API_KEY1, model="llama3-8b-8192")
        self.tools = []

    def research(self, extracted_instructions):
        try:
            # Use the prompt to generate chapter content
            research_agent = create_tool_calling_agent(
                self.llm, self.tools, self.prompt
            )
            research_agent_executor = AgentExecutor(
                agent=research_agent, tools=self.tools
            )
            response = research_agent_executor.invoke({"topic": extracted_instructions.get("title"), "chapters": extracted_instructions.get("chapters"), "writing_style": extracted_instructions.get("writing_style")})
            response = response.get("output")

            chapters = ast.literal_eval(response)

            return chapters

        except Exception as e:
            logger.exception(f"Error: {e} | Traceback: {traceback.format_exc()}")
            raise e


class WriterAgent:
    def __init__(self):
        self.prompt = PromptTemplate(
            input_variables=["topic", "chapter" "writing_style"],
            template="""You are an expert eBook writer. Your task is to write a detailed, well-structured, and engaging chapter for an eBook based on the given chapter title and overall topic of the book. The chapter should be informative, easy to understand, and provide in-depth coverage of the subject matter.

        Here are your instructions:

        1. The content should be structured with proper headings, subheadings, and sections.
        2. The writing should be in a clear and engaging style, suitable for a broad audience such as students, researchers, or curious readers.
        3. Ensure the chapter content stays focused on the chapter title, but it should also be aligned with the overall topic of the book.
        4. Provide relevant examples, explanations, and any necessary background information to make the content more comprehensive and easier to understand.
        5. The length of the chapter should be around 1500-2000 words.
        6. Use a formal yet approachable tone. Avoid overly technical jargon unless it's explained well.
        7. Writing style should be {writing_style}.

        Inputs:
        - **Topic**: {topic}
        - **Chapter Title**: {chapter}

        Now, based on the provided inputs, write the complete content for this chapter. The content should be returned as a single string of text.

        Note: Don't include chapter number and chapter title in the content. Also don't include your instructions in the content.
        \n\n placeholder:{agent_scratchpad}
        """,
        )

        self.llm = ChatGroq(api_key=Config.GROQ_API_KEY2, model="Llama-3.1-8b-instant")
        self.tools = []

    def write(self, chapter, extracted_instructions):
        try:
            # Use the prompt to generate chapter titles
            research_agent = create_tool_calling_agent(
                self.llm, self.tools, self.prompt
            )
            research_agent_executor = AgentExecutor(
                agent=research_agent, tools=self.tools
            )
            response = research_agent_executor.invoke(
                {"topic": extracted_instructions.get("title"), "chapter": chapter, "writing_style": extracted_instructions.get("writing_style")}
            )
            chapter_content = response.get("output")

            return chapter_content

        except Exception as e:
            logger.exception(f"Error: {e} | Traceback: {traceback.format_exc()}")
            raise e


def generate_ebook(topic, chapters_content, chapters, output_format="pdf"):
    try:
        # Ensure output directory exists
        os.makedirs("ebooks", exist_ok=True)

        # Generate filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"ebooks/{topic.replace(' ', '_')}_{timestamp}"

        if output_format.lower() == "pdf":
            pdf = HTML2PDF("P", "mm", "A4")
            pdf.set_auto_page_break(auto=True, margin=15)

            # Cover Page with improved styling
            pdf.add_page()
            pdf.set_font(family="Helvetica", style="B", size=30)
            pdf.set_text_color(0, 0, 0)
            pdf.set_y(pdf.h / 2 - 15)
            
            # Centered title with subtle background
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(
                w=pdf.w - 20, h=20, text=topic, align="C", 
                fill=True, new_x="LMARGIN", new_y="NEXT"
            )

            # Table of Contents with improved layout
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(0, 10, text="Table of Contents", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(5)
            
            pdf.set_font("Helvetica", size=12)
            pdf.set_text_color(50, 50, 50)
            for i, chapter_title in enumerate(chapters):
                # Add dots between chapter title and page number
                pdf.cell(
                    0, 10,
                    text=f"Chapter {i+1}: {chapter_title} " + "." * 50,
                    new_x="LMARGIN", 
                    new_y="NEXT"
                )

            # Chapters
            for i, content in enumerate(chapters_content):
                pdf.add_page()
                pdf.chapter_title(i+1, chapters[i])
                pdf.chapter_body(content)

            # Save PDF file
            pdf_file_path = f"{base_filename}.pdf"
            pdf.output(pdf_file_path)
            return pdf_file_path

        elif output_format.lower() == "doc":
            # Create Word document
            doc = Document()
            doc.add_heading(topic, level=1)

            # Table of Contents
            doc.add_heading("Table of Contents", level=2)
            for i, chapter_title in enumerate(chapters):
                doc.add_paragraph(f"Chapter {i+1}: {chapter_title}")

            # Chapters
            for i, content in enumerate(chapters_content):
                doc.add_heading(f"Chapter {i+1}: {chapters[i]}", level=2)
                doc.add_paragraph(content)

            # Save Word document
            doc_file_path = f"{base_filename}.docx"
            doc.save(doc_file_path)
            return doc_file_path

        elif output_format.lower() == "text":
            # Create TXT file
            txt_file_path = f"{base_filename}.txt"
            with open(txt_file_path, "w", encoding="utf-8") as txt_file:
                # Write title
                txt_file.write(f"{topic}\n\n")

                # Write Table of Contents
                txt_file.write("Table of Contents\n")
                for i, chapter_title in enumerate(chapters):
                    txt_file.write(f"Chapter {i+1}: {chapter_title}\n")
                txt_file.write("\n")

                # Write Chapters
                for i, content in enumerate(chapters_content):
                    txt_file.write(f"Chapter {i+1}: {chapters[i]}\n")
                    txt_file.write(f"{content}\n\n")

            return txt_file_path

        else:
            raise ValueError(f"Unsupported output format: {output_format}")

    except Exception as e:
        logger.exception(f"Error: {e} | Traceback: {traceback.format_exc()}")
        raise e


class EbookGenerator:
    @staticmethod
    def generate_ebook_task(extracted_instructions):
        try:
            if not extracted_instructions:
                raise Exception("Data not provided")

            # Initialize agents
            logger.info("Initializing agents...")
            researcher = ResearcherAgent()
            writer = WriterAgent()
            editor = EditorAgent()

            # Step 1: Research the topic and get chapter titles
            logger.info("Researching the topic...")
            chapters = researcher.research(extracted_instructions)

            # Step 2: Write content for each chapter
            logger.info("Writing content for each chapter...")
            written_content = [writer.write(chapter, extracted_instructions) for chapter in chapters]

            # Step 3: Edit each chapter for grammar and consistency
            logger.info("Editing each chapter for grammar and consistency...")
            final_content = [editor.edit(content, extracted_instructions.get("writing_style")) for content in written_content]

            # Step 4: Decoding text to latin-1
            logger.info("Decoding text to latin-1...")
            final_content = [
                content.encode("latin-1", "ignore").decode("latin-1")
                for content in final_content
            ]

            # Step 5: Generate PDF with all content
            logger.info("Generating PDF with all content...")
            file_path = generate_ebook(extracted_instructions.get("title"), final_content, chapters, extracted_instructions.get("format"))

            # Return PDF file path
            logger.info("Task completed.")

            collected = gc.collect()

            logger.info(f"Garbage collector: collected {collected} objects.")
            return file_path

        except Exception as e:
            logger.exception(f"Error: {e} | Traceback: {traceback.format_exc()}")
            raise e


def reply_with_bot(mention: dict):
    """Reply to a mention with a PDF eBook."""
    try:
        logger.info("Processing mention reply_with_pdf ...")
        # Extract topic from the mention text
        text = mention.record.text
        words = text.split()
        mentioned_text = " ".join(words[1:])  # Assuming the topic is after the bot mention
        suggested_text = None
        public_id = None

        result = extract_ebook_instructions(mentioned_text)
        extracted_instructions = json.loads(result.model_dump_json())
        logger.info(f"Topic: {extracted_instructions.get('title')}")
        
        if extracted_instructions.get("is_suggestion") == 1:

            # Initialize Groq LLM
            llm = ChatGroq(api_key=os.getenv("GROQ_API_KEY1"), model_name="mixtral-8x7b-32768")

            # Create a prompt template
            prompt_template = PromptTemplate(
                input_variables=["question"],
                template="Drawing from the depths of human knowledge and wisdom, craft a thoughtful, illuminating response to the following inquiry that captures its essence with eloquence and brevity: {question} Distill the profound into a spark of insight that ignites understanding. \n Note: Provide a concise and to-the-point answer to the following question in no more than 280 characters."
            )

            # Create an LLM chain
            chain = LLMChain(llm=llm, prompt=prompt_template)

            # Generate response
            suggested_text = chain.run(mentioned_text)

            # Ensure response is within character limit
            if len(suggested_text) > 280:
                # If too long, split and take the first part
                splitter = CharacterTextSplitter(
                    separator="\n",
                    chunk_size=280,
                    chunk_overlap=0,
                    length_function=len,
                )
                chunks = splitter.split_text(suggested_text)
                suggested_text = chunks[0]
        else:
            # Generate content and create ebook file

            file_path = EbookGenerator.generate_ebook_task(extracted_instructions)
            logger.info(f"file path: {file_path}")

            # Upload PDF to Cloudinary
            download_url, public_id = upload_to_cloudinary(file_path, extracted_instructions.get('title'))
            
            logger.info(f"PDF uploaded to: {download_url}")

        try:
            # Create proper reply reference
            reply_ref = {
                "root": {"cid": mention.cid, "uri": mention.uri},
                "parent": {"cid": mention.cid, "uri": mention.uri},
            }

            if not suggested_text:
                # Create the embed external object with the download link
                embed = models.AppBskyEmbedExternal.Main(
                    external=models.AppBskyEmbedExternal.External(
                        title=f"Ebook: {extracted_instructions.get('title')}",
                        description="Click to download your generated ebook",
                        uri=download_url,
                        thumb=None,
                    )
                )

                # Reply to the mention with the download link
                client.post(
                    text=f"🤖 Here's your ebook about {extracted_instructions.get('title')}! 📚\nClick here to download: {download_url}",
                    reply_to=reply_ref,
                    embed=embed,
                )

                logger.info(f"PDF link sent successfully!")

                # Clean up local file
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.info(f"Local PDF file cleaned up: {file_path}")

            else:
                # Reply to the mention with the download link
                client.post(
                    text=f"{suggested_text}",
                    reply_to=reply_ref,
                )

        except Exception as e:
            # If posting fails, clean up the uploaded file from Cloudinary
            delete_from_cloudinary(public_id)
            raise e

    except Exception as e:
        logger.exception(f"Error in reply_with_pdf: {str(e)}")
        print(f"Error: {str(e)}")


"""
Bluesky Book Suggestion thread Bot
"""
class BlueskyBookSuggestionBot:
    def __init__(self, post_interval_minutes=30):
        # Initialize Bluesky client
        self.bluesky_client = client
        
        # Initialize Groq LLM
        self.llm = ChatGroq(
            temperature=0.7, 
            model_name="llama-3.3-70b-versatile",
            groq_api_key=os.getenv('GROQ_API_KEY1')
        )
        
        # Scheduling parameters
        self.post_interval = post_interval_minutes
        
        # Tracking engagement
        self.post_engagement_history = []
        
    def generate_book_topic(self):
        """Generate an engaging book-related topic."""
        topic_prompt = PromptTemplate(
            input_variables=['genre'],
            template="""
            You are a book suggestion twitter bot. Generate a provocative, boundary-pushing book discussion topic in the {genre} genre. 
            Create a thread that:
            - Challenges conventional wisdom
            - Reveals uncomfortable truths
            - Sparks intense intellectual debate
            - Uses a raw, unfiltered, and unapologetic tone
            - Combines intellectual depth with viral potential

            RULES:
            - Be brutally honest
            - Use sharp, confrontational language
            - Expose hidden narratives
            - Craft a hook that demands attention
            - Make people uncomfortable enough to engage
            - Synthesize complex ideas into razor-sharp statements
            - Each thread part should be a punch to intellectual complacency

            Tone: Uncompromising. Cerebral. Confrontational.
            Goal: Intellectual provocation that breaks echo chambers

            **Output Format Instructions**:  
            - The output must be a valid Python list that contains exactly 10 thread parts as comma separated strings.
            - Each thread part should be a properly formatted python string without any special keywords, symbols, contractions or unnecessary punctuation.
            - Ensure threads are meaningful and follow a logical progression.

            **output Format must be like below list example**:
            ['thread part 1', 'thread part 2', 'thread part 3', 'thread part 4', 'thread part 5', 'thread part 6', 'thread part 7', 'thread part 8', 'thread part 9', 'thread part 10']

            Note: don't include any explanation from your side.
            """
        )
        
        # Rotate through genres to keep content diverse
        genres = [
            # Technology
            'technology', 
            'artificial intelligence', 
            'cybersecurity', 
            'computer science', 
            'digital innovation',
            'emerging technologies',

            # Science
            'science', 
            'astronomy', 
            'biology', 
            'physics', 
            'environmental science', 
            'neuroscience',
            'quantum physics',
            'climate science',

            # Mathematics
            'mathematics', 
            'applied mathematics', 
            'data science', 
            'cryptography', 
            'statistical analysis',
            'computational mathematics',

            # Philosophy
            'philosophy', 
            'ethics', 
            'political philosophy', 
            'existentialism', 
            'epistemology',
            'philosophy of science',
            'logic',

            # History
            'history', 
            'world history', 
            'military history', 
            'cultural history', 
            'ancient civilizations',
            'modern history',
            'diplomatic history',
            'social movements',

            # Autobiography
            'autobiography', 
            'scientific biography', 
            'tech innovators', 
            'political leadership', 
            'explorers and pioneers',
            'intellectual memoirs',
            'social activists'
        ]
        
        # Create chain
        topic_chain = topic_prompt | self.llm
        genre = random.choice(genres)
        
        # Generate topic
        topic_response = topic_chain.invoke({"genre": genre})
        return topic_response
    
    def create_threaded_post(self, main_topic):
        """Create a threaded post with multiple parts."""
        try:
            # Parse the topic into a list of thread parts
            topic_parts = ast.literal_eval(main_topic.content)
            
            if not isinstance(topic_parts, list) or len(topic_parts) < 1:
                raise ValueError("Generated topic must be a list with at least one part.")
            
            # Post the root thread (first part)
            root_post = self.bluesky_client.send_post(text=topic_parts[0])
            
            # Initialize parent references for threading
            parent_uri = root_post.uri
            parent_cid = root_post.cid
            
            # Iterate over the remaining parts and post them as replies
            for part in topic_parts[1:]:
                if part.strip():  # Skip empty parts
                    thread_post = self.bluesky_client.send_post(
                        text=part,
                        reply_to={
                            'root': {'uri': root_post.uri, 'cid': root_post.cid},
                            'parent': {'uri': parent_uri, 'cid': parent_cid}
                        }
                    )
                    
                    # Update parent references for the next post in the thread
                    parent_uri = thread_post.uri
                    parent_cid = thread_post.cid
            
            return {"root_post": root_post, "thread_posts": []}
    
        except Exception as e:
            print(f"Error creating threaded post: {e}")
            return None
    
    def run_post_cycle(self):
        """
        Single post cycle for the bot.
        Generates topic, creates post, and tracks engagement.
        """
        try:
            # Generate book topic
            book_topic = self.generate_book_topic()
            
            # Create and post threaded content
            posted_thread = self.create_threaded_post(book_topic)
            
            if posted_thread:
                # Track engagement (simulated for this example)
                self.post_engagement_history.append({
                    'likes': random.randint(10, 100),
                    'post': posted_thread
                })
                
                print(f"Successfully posted book discussion thread at {datetime.datetime.now()}")
            
        except Exception as e:
            print(f"Error in post cycle: {e}")
    
    def start(self):
        """Start the bot's posting schedule."""
        # Schedule the post cycle
        schedule.every(self.post_interval).minutes.do(self.run_post_cycle)
        
        print(f"Book Suggestion Bot started. Posting every {self.post_interval} minutes.")
        
    
    def stop(self):
        """Stops all scheduled jobs."""
        schedule.clear()
        print("Book Suggestion Bot stopped.")

def main():
    """Main bot loop to monitor and reply to mentions."""
    try:
        print("Bot started. Listening for mentions...")
        # Keep track of processed notifications
        processed_notifications = set()
        last_processed_time = datetime.datetime.now(datetime.timezone.utc)

        # Initialize the bot with 30-minute intervals
        bot = BlueskyBookSuggestionBot(post_interval_minutes=30)
        # Start the bot
        bot.start()

        while True:
            
            notifications = client.app.bsky.notification.list_notifications().notifications

            for note in notifications:
                if (
                    note.reason == "mention"
                    and note.uri not in processed_notifications
                    and datetime.datetime.strptime(
                        note.indexed_at, "%Y-%m-%dT%H:%M:%S.%fZ"
                    ).replace(tzinfo=datetime.timezone.utc)
                    > last_processed_time
                ):
                    reply_with_bot(note)
                    processed_notifications.add(note.uri)

            # Limit processed notifications to prevent memory growth
            if len(processed_notifications) > 1000:
                processed_notifications = set(list(processed_notifications)[-500:])

            schedule.run_pending()

            time.sleep(3)  # Adjust the sleep duration as needed

    except KeyboardInterrupt as e:
        bot.stop()
        logger.exception(f"Error in main (): {e} | Traceback: {traceback.format_exc()}")
    except Exception as e:
        bot.stop()
        logger.exception(f"Error in main (): {e} | Traceback: {traceback.format_exc()}")

if __name__ == "__main__":
    main()
